"""
docx_parser.py - Word document (DOCX) text extraction module for LexiVault.
"""
import docx

def parse_docx(file_path: str) -> str:
    """Parse a DOCX file and extract clean text.

    Args:
        file_path: Absolute or relative path to the DOCX file.

    Returns:
        Full clean text string. Returns an empty string if parsing fails.
    """
    try:
        doc = docx.Document(file_path)
        full_text_parts = []
        
        # Extract paragraph texts
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text_parts.append(text)
                
        # Extract table texts
        for table in doc.tables:
            for row in table.rows:
                # Deduplicate cells to avoid repeated text in merged cells
                row_cells = []
                last_cell_text = None
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text != last_cell_text:
                        row_cells.append(cell_text)
                        last_cell_text = cell_text
                if row_cells:
                    full_text_parts.append(" | ".join(row_cells))
                    
        return "\n\n".join(full_text_parts)
    except Exception as e:
        print(f"Error parsing DOCX '{file_path}': {str(e)}")
        return ""
