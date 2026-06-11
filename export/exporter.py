"""
exporter.py - Report export module for LexiVault.

Exports decision briefs and risk analysis results as professionally formatted
PDF (via ReportLab) and Word (via python-docx) documents with color-coded
risk levels and LexiVault branding.
"""

import os
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Color mapping for risk levels
RISK_COLORS_RGB = {
    "High": (255, 0, 0),       # Red
    "Medium": (255, 165, 0),   # Orange
    "Low": (0, 128, 0),        # Green
}

RISK_COLORS_REPORTLAB = {
    "High": colors.red,
    "Medium": colors.orange,
    "Low": colors.green,
}


def export_pdf(decision_brief: str, risk_scores: list, filename: str) -> str:
    """Export decision brief and risk scores as a professionally formatted PDF.

    Generates a PDF with LexiVault branding, including a header with product
    name and tagline, executive summary with the decision brief text, a
    color-coded risk dashboard table, and a branded footer.

    Args:
        decision_brief: Full decision brief text string to include in the
                        Executive Summary section.
        risk_scores: List of dicts, each with keys 'clause_name', 'value',
                     and 'risk_level' (High/Medium/Low).
        filename: Output file path string where the PDF will be saved.

    Returns:
        The saved file path string. Returns an empty string if export fails.
    """
    try:
        # Ensure output directory exists
        output_dir = os.path.dirname(filename)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        styles = getSampleStyleSheet()
        elements = []

        # Custom styles
        header_style = ParagraphStyle(
            "LexiVaultHeader",
            parent=styles["Title"],
            fontSize=28,
            textColor=colors.HexColor("#1a1a2e"),
            spaceAfter=6,
            alignment=1,  # Center
        )

        tagline_style = ParagraphStyle(
            "Tagline",
            parent=styles["Normal"],
            fontSize=14,
            textColor=colors.HexColor("#4a4a8a"),
            spaceAfter=24,
            alignment=1,  # Center
            fontName="Helvetica-Oblique",
        )

        section_style = ParagraphStyle(
            "SectionTitle",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=colors.HexColor("#1a1a2e"),
            spaceBefore=18,
            spaceAfter=12,
        )

        body_style = ParagraphStyle(
            "BodyText",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            spaceAfter=8,
        )

        footer_style = ParagraphStyle(
            "Footer",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.grey,
            alignment=1,  # Center
            spaceBefore=36,
        )

        # ---- Header ----
        elements.append(Paragraph("🏛️ LexiVault", header_style))
        elements.append(Paragraph("Stop Hunting. Start Deciding.", tagline_style))
        elements.append(Spacer(1, 12))

        # ---- Executive Summary ----
        elements.append(Paragraph("Executive Summary", section_style))

        # Split brief into paragraphs for better formatting
        brief_paragraphs = decision_brief.split("\n")
        for para in brief_paragraphs:
            para = para.strip()
            if para:
                # Escape special characters for ReportLab
                safe_para = (
                    para.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                elements.append(Paragraph(safe_para, body_style))

        elements.append(Spacer(1, 18))

        # ---- Risk Dashboard ----
        elements.append(Paragraph("Risk Dashboard", section_style))

        if risk_scores:
            # Build table data
            table_data = [["Clause", "Risk Level", "Value"]]

            for risk in risk_scores:
                clause_name = str(risk.get("clause_name", ""))
                risk_level = str(risk.get("risk_level", ""))
                value = str(risk.get("value", ""))

                # Truncate long values for table readability
                if len(value) > 80:
                    value = value[:77] + "..."

                table_data.append([clause_name, risk_level, value])

            table = Table(table_data, colWidths=[1.8 * inch, 1.2 * inch, 3.5 * inch])

            # Build table style with color-coded risk rows
            style_commands = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 11),
                ("FONTSIZE", (0, 1), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("TOPPADDING", (0, 1), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 1), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]

            # Color-code each row based on risk level
            for i, risk in enumerate(risk_scores, start=1):
                risk_level = str(risk.get("risk_level", ""))
                color = RISK_COLORS_REPORTLAB.get(risk_level, colors.black)
                style_commands.append(("TEXTCOLOR", (0, i), (-1, i), color))

                # Light background tint for alternating rows
                if i % 2 == 0:
                    style_commands.append(
                        ("BACKGROUND", (0, i), (-1, i), colors.HexColor("#f5f5f5"))
                    )

            table.setStyle(TableStyle(style_commands))
            elements.append(table)
        else:
            elements.append(Paragraph("No risk analysis data available.", body_style))

        # ---- Footer ----
        elements.append(Spacer(1, 36))
        elements.append(
            Paragraph(
                "Generated by LexiVault — Private AI for Document Decisions",
                footer_style,
            )
        )

        # Build the PDF
        doc.build(elements)
        return filename

    except Exception as e:
        print(f"Error exporting PDF: {str(e)}")
        return ""


def export_docx(decision_brief: str, risk_scores: list, filename: str) -> str:
    """Export decision brief and risk scores as a professionally formatted Word document.

    Generates a DOCX file with LexiVault branding, including a header with product
    name and tagline, executive summary with the decision brief text, a
    color-coded risk dashboard table, and a branded footer.

    Args:
        decision_brief: Full decision brief text string to include in the
                        Executive Summary section.
        risk_scores: List of dicts, each with keys 'clause_name', 'value',
                     and 'risk_level' (High/Medium/Low).
        filename: Output file path string where the DOCX will be saved.

    Returns:
        The saved file path string. Returns an empty string if export fails.
    """
    try:
        # Ensure output directory exists
        output_dir = os.path.dirname(filename)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc = Document()

        # ---- Header ----
        title = doc.add_heading("LexiVault", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(26, 26, 46)

        tagline = doc.add_paragraph("Stop Hunting. Start Deciding.")
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in tagline.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(74, 74, 138)
            run.font.italic = True

        doc.add_paragraph()  # Spacer

        # ---- Executive Summary ----
        doc.add_heading("Executive Summary", level=1)

        brief_paragraphs = decision_brief.split("\n")
        for para in brief_paragraphs:
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        doc.add_paragraph()  # Spacer

        # ---- Risk Dashboard ----
        doc.add_heading("Risk Dashboard", level=1)

        if risk_scores:
            # Create table with header row
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            # Header cells
            header_cells = table.rows[0].cells
            header_cells[0].text = "Clause"
            header_cells[1].text = "Risk Level"
            header_cells[2].text = "Value"

            # Style header cells
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                from docx.oxml.ns import qn
                shading_elm = cell._element.get_or_add_tcPr()
                shading = shading_elm.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:fill"): "1a1a2e",
                        qn("w:val"): "clear",
                    },
                )
                shading_elm.append(shading)

            # Data rows
            for risk in risk_scores:
                clause_name = str(risk.get("clause_name", ""))
                risk_level = str(risk.get("risk_level", ""))
                value = str(risk.get("value", ""))

                row_cells = table.add_row().cells
                row_cells[0].text = clause_name
                row_cells[1].text = risk_level
                row_cells[2].text = value

                # Color-code based on risk level
                rgb = RISK_COLORS_RGB.get(risk_level, (0, 0, 0))
                color = RGBColor(*rgb)

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = color
                            run.font.size = Pt(10)

            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(2.0)
                row.cells[1].width = Inches(1.5)
                row.cells[2].width = Inches(3.5)
        else:
            doc.add_paragraph("No risk analysis data available.")

        doc.add_paragraph()  # Spacer

        # ---- Footer ----
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            "Generated by LexiVault — Private AI for Document Decisions"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_run.font.italic = True

        # Save the document
        doc.save(filename)
        return filename

    except Exception as e:
        print(f"Error exporting DOCX: {str(e)}")
        return ""
